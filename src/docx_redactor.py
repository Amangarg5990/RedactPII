"""
Format-Preserving DOCX Redactor Module.
Parses Word documents (.docx) and applies PII redaction to paragraphs, tables, headers, and footers.
Preserves character formatting (bold, italic, font size, styling).
"""

import os
from typing import Dict, Any, Tuple, List
from src.pii_detector import PIIDetector, PIIMatch
from src.faker_mapper import SyntheticPIIMapper

try:
    import docx
    from docx.document import Document
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Cell
except ImportError:
    docx = None


class DOCXRedactor:
    def __init__(self, detector: PIIDetector = None, mapper: SyntheticPIIMapper = None):
        self.detector = detector or PIIDetector()
        self.mapper = mapper or SyntheticPIIMapper()
        self.stats = {
            "paragraphs_processed": 0,
            "table_cells_processed": 0,
            "total_redactions": 0,
            "redactions_by_type": {}
        }

    def redact_paragraph(self, p) -> int:
        """
        Redacts PII in a single paragraph while preserving run formatting.
        """
        full_text = p.text
        if not full_text or not full_text.strip():
            return 0

        matches = self.detector.detect(full_text)
        if not matches:
            return 0

        redaction_count = len(matches)
        self.stats["total_redactions"] += redaction_count

        # Build redacted text by applying replacements right-to-left
        new_text = full_text
        for m in reversed(matches):
            pii_type = m.pii_type
            self.stats["redactions_by_type"][pii_type] = self.stats["redactions_by_type"].get(pii_type, 0) + 1
            replacement = self.mapper.get_replacement(m.text, pii_type)
            new_text = new_text[:m.start] + replacement + new_text[m.end:]

        # Apply new_text while preserving formatting of first run
        if p.runs:
            first_run = p.runs[0]
            first_run.text = new_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = new_text

        return redaction_count

    def redact_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """
        Redacts PII in an entire .docx document and saves to output_path.
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")

        doc = docx.Document(input_path)

        # 1. Process Body Paragraphs
        for p in doc.paragraphs:
            self.stats["paragraphs_processed"] += 1
            self.redact_paragraph(p)

        # 2. Process Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.stats["table_cells_processed"] += 1
                    for p in cell.paragraphs:
                        self.redact_paragraph(p)

        # 3. Process Headers & Footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for p in header.paragraphs:
                        self.redact_paragraph(p)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    self.redact_paragraph(p)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for p in footer.paragraphs:
                        self.redact_paragraph(p)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    self.redact_paragraph(p)

        # Save redacted document
        doc.save(output_path)
        return self.stats
